#!/usr/bin/env python3
"""
Extraction de texte depuis un fichier DOCX.
Utilise python-docx pour extraire le texte structuré par paragraphes.
"""

import argparse
import json
import os
import sys

try:
    from docx import Document
except ImportError:
    print("Erreur : python-docx requis. Installez avec : pip install python-docx")
    sys.exit(1)


def extract_docx(filepath: str, output_dir: str) -> None:
    """
    Extrait le texte d'un fichier DOCX avec détection des titres.
    
    Args:
        filepath: Chemin vers le fichier DOCX
        output_dir: Dossier de sortie pour les fichiers JSON
    """
    doc = Document(filepath)

    # Métadonnées
    metadata = {
        'title': doc.core_properties.title or '',
        'author': doc.core_properties.author or ''
    }

    with open(os.path.join(output_dir, 'metadata.json'), 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

    # Extraction des chapitres basée sur les styles de titre
    chapters = []
    current_chapter = None
    current_text = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else 'Normal'

        # Détection des titres (Heading 1, Heading 2, Titre 1, etc.)
        if 'heading' in style_name.lower() or 'titre' in style_name.lower() or 'title' in style_name.lower():
            # Sauvegarder le chapitre précédent
            if current_chapter is not None and current_text:
                chapters.append({
                    'title': current_chapter,
                    'text': '\n\n'.join(current_text)
                })

            current_chapter = text
            current_text = []
        else:
            current_text.append(text)

    # Dernier chapitre
    if current_chapter is not None and current_text:
        chapters.append({
            'title': current_chapter,
            'text': '\n\n'.join(current_text)
        })

    # Si aucun chapitre détecté, tout mettre dans un seul chapitre
    if not chapters:
        all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        chapters.append({
            'title': 'Document',
            'text': '\n\n'.join(all_text)
        })

    with open(os.path.join(output_dir, 'chapters.json'), 'w', encoding='utf-8') as f:
        json.dump(chapters, f, ensure_ascii=False, indent=2)

    print(f"Extraction terminée : {len(chapters)} chapitres extraits")
    print(f"Métadonnées : {metadata['title']} par {metadata['author']}")


def main():
    parser = argparse.ArgumentParser(description='Extraction de texte depuis DOCX')
    parser.add_argument('--input', required=True, help='Chemin vers le fichier DOCX')
    parser.add_argument('--output', required=True, help='Dossier de sortie')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Erreur : fichier introuvable : {args.input}")
        sys.exit(1)

    os.makedirs(args.output, exist_ok=True)
    extract_docx(args.input, args.output)


if __name__ == '__main__':
    main()
